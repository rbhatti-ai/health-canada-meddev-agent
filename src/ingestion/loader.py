"""
Document loader for various file formats.

Supports:
- PDF files (using PyMuPDF)
- DOCX files (using python-docx)
- Markdown files
- Plain text files
"""

import hashlib
from collections.abc import Generator
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from src.utils.logging import get_logger

logger = get_logger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document."""

    content: str
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def id(self) -> str:
        """Generate unique ID based on content hash."""
        return hashlib.md5(self.content.encode()).hexdigest()[:16]


@dataclass
class LoadedDocument:
    """Represents a fully loaded document."""

    source_path: str
    file_type: str
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    pages: list[str] = field(default_factory=list)

    @property
    def id(self) -> str:
        """Generate unique ID based on source path."""
        return hashlib.md5(self.source_path.encode()).hexdigest()[:16]

    @property
    def page_count(self) -> int:
        """Number of pages (if applicable)."""
        return len(self.pages) if self.pages else 1


class DocumentLoader:
    """
    Multi-format document loader for regulatory documents.

    Supports PDF, DOCX, MD, and TXT files with metadata extraction.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".md", ".txt", ".markdown"}

    def __init__(self) -> None:
        self.logger = get_logger(self.__class__.__name__)

    def load_file(self, file_path: Path) -> LoadedDocument | None:
        """
        Load a single document file.

        Args:
            file_path: Path to the document

        Returns:
            LoadedDocument or None if loading fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            return None

        extension = file_path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            self.logger.warning(f"Unsupported file type: {extension}")
            return None

        try:
            if extension == ".pdf":
                return self._load_pdf(file_path)
            elif extension in {".docx", ".doc"}:
                return self._load_docx(file_path)
            elif extension in {".md", ".markdown"}:
                return self._load_markdown(file_path)
            elif extension == ".txt":
                return self._load_text(file_path)
        except Exception as e:
            self.logger.error(f"Error loading {file_path}: {e}")
            return None

        return None

    def load_directory(
        self,
        directory: Path,
        recursive: bool = True,
    ) -> Generator[LoadedDocument, None, None]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Path to directory
            recursive: Whether to process subdirectories

        Yields:
            LoadedDocument for each successfully loaded file
        """
        directory = Path(directory)

        if not directory.is_dir():
            self.logger.error(f"Not a directory: {directory}")
            return

        pattern = "**/*" if recursive else "*"

        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                doc = self.load_file(file_path)
                if doc:
                    yield doc

    def _load_pdf(self, file_path: Path) -> LoadedDocument:
        """Load a PDF file using PyMuPDF."""
        self.logger.info(f"Loading PDF: {file_path.name}")

        doc = fitz.open(file_path)
        pages = []
        full_text_parts = []

        for _page_num, page in enumerate(doc):
            text = page.get_text("text")
            pages.append(text)
            full_text_parts.append(text)

        full_text = "\n\n".join(full_text_parts)

        # Extract metadata
        metadata = {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": "pdf",
            "page_count": len(pages),
            "title": doc.metadata.get("title", file_path.stem),
            "author": doc.metadata.get("author", ""),
            "creation_date": doc.metadata.get("creationDate", ""),
            "modification_date": doc.metadata.get("modDate", ""),
            "loaded_at": datetime.now().isoformat(),
        }

        # Try to detect document category from path or content
        metadata["category"] = self._detect_category(file_path, full_text)

        doc.close()

        return LoadedDocument(
            source_path=str(file_path),
            file_type="pdf",
            content=full_text,
            metadata=metadata,
            pages=pages,
        )

    def _load_docx(self, file_path: Path) -> LoadedDocument:
        """Load a DOCX file using python-docx."""
        self.logger.info(f"Loading DOCX: {file_path.name}")

        doc = DocxDocument(file_path)

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)

        # Extract core properties
        core_props = doc.core_properties

        metadata = {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": "docx",
            "title": core_props.title or file_path.stem,
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "loaded_at": datetime.now().isoformat(),
            "category": self._detect_category(file_path, full_text),
        }

        return LoadedDocument(
            source_path=str(file_path),
            file_type="docx",
            content=full_text,
            metadata=metadata,
        )

    def _load_markdown(self, file_path: Path) -> LoadedDocument:
        """Load a Markdown file."""
        self.logger.info(f"Loading Markdown: {file_path.name}")

        content = file_path.read_text(encoding="utf-8")

        # Extract title from first heading if present
        title = file_path.stem
        lines = content.split("\n")
        for line in lines:
            if line.startswith("# "):
                title = line[2:].strip()
                break

        metadata = {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": "markdown",
            "title": title,
            "loaded_at": datetime.now().isoformat(),
            "category": self._detect_category(file_path, content),
        }

        return LoadedDocument(
            source_path=str(file_path),
            file_type="markdown",
            content=content,
            metadata=metadata,
        )

    def _load_text(self, file_path: Path) -> LoadedDocument:
        """Load a plain text file."""
        self.logger.info(f"Loading text file: {file_path.name}")

        content = file_path.read_text(encoding="utf-8")

        metadata = {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": "text",
            "title": file_path.stem,
            "loaded_at": datetime.now().isoformat(),
            "category": self._detect_category(file_path, content),
        }

        return LoadedDocument(
            source_path=str(file_path),
            file_type="text",
            content=content,
            metadata=metadata,
        )

    def _detect_category(self, file_path: Path, content: str) -> str:
        """
        Detect document category based on path and content.

        Categories:
        - regulation: Core laws and regulations
        - guidance: Health Canada guidance documents
        - standard: ISO and other standards
        - form: Application forms
        - checklist: Checklists and summaries
        - other: Uncategorized
        """
        path_lower = str(file_path).lower()
        content_lower = content[:2000].lower()  # Check first 2000 chars

        # Check path patterns
        if "regulation" in path_lower or "law" in path_lower:
            return "regulation"
        if "guidance" in path_lower:
            return "guidance"
        if "standard" in path_lower or "iso" in path_lower:
            return "standard"
        if "form" in path_lower:
            return "form"
        if "checklist" in path_lower or "summary" in path_lower:
            return "checklist"

        # Check content patterns
        if "sor/98-282" in content_lower or "medical devices regulations" in content_lower:
            return "regulation"
        if "guidance document" in content_lower:
            return "guidance"
        if "iso 13485" in content_lower or "iec 62304" in content_lower:
            return "standard"

        return "other"


# Singleton instance
document_loader = DocumentLoader()
