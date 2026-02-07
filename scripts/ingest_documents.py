#!/usr/bin/env python3
"""
Standalone document ingestion script.
Ingests regulatory documents into ChromaDB vector store.
"""

import hashlib
import re
import sys
from pathlib import Path
from typing import Any

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent))

import chromadb
import fitz  # PyMuPDF
from chromadb.config import Settings as ChromaSettings
from docx import Document as DocxDocument
from openai import OpenAI

# Configuration
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
EMBEDDING_MODEL = "text-embedding-3-small"
COLLECTION_NAME = "health_canada_regulatory"


def load_pdf(file_path: Path) -> dict[str, Any]:
    """Load a PDF file."""
    print(f"  Loading PDF: {file_path.name}")
    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        text = page.get_text("text")
        pages.append(text)
    full_text = "\n\n".join(pages)
    doc.close()

    return {
        "content": full_text,
        "metadata": {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": "pdf",
            "page_count": len(pages),
        },
    }


def load_docx(file_path: Path) -> dict[str, Any]:
    """Load a DOCX file."""
    print(f"  Loading DOCX: {file_path.name}")
    doc = DocxDocument(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return {
        "content": "\n\n".join(paragraphs),
        "metadata": {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": "docx",
        },
    }


def load_markdown(file_path: Path) -> dict[str, Any]:
    """Load a Markdown file."""
    print(f"  Loading MD: {file_path.name}")
    content = file_path.read_text(encoding="utf-8")
    return {
        "content": content,
        "metadata": {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": "markdown",
        },
    }


def detect_category(file_path: Path, content: str) -> str:
    """Detect document category."""
    path_lower = str(file_path).lower()
    content_lower = content[:2000].lower()

    if "regulation" in path_lower or "law" in path_lower:
        return "regulation"
    if "guidance" in path_lower:
        return "guidance"
    if "standard" in path_lower or "iso" in path_lower:
        return "standard"
    if "form" in path_lower:
        return "form"
    if "checklist" in path_lower or "summary" in path_lower:
        return "checklist"
    if "sor/98-282" in content_lower or "medical devices regulations" in content_lower:
        return "regulation"
    if "guidance document" in content_lower:
        return "guidance"
    if "iso 13485" in content_lower or "iec 62304" in content_lower:
        return "standard"

    return "other"


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into chunks with overlap."""
    # Normalize whitespace
    text = re.sub(r"\s+", " ", text).strip()

    if len(text) <= chunk_size:
        return [text] if len(text) > 50 else []

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence end in last 20% of chunk
            search_start = start + int(chunk_size * 0.8)
            search_region = text[search_start:end]

            # Find last sentence boundary
            for sep in [". ", "! ", "? ", "\n"]:
                last_sep = search_region.rfind(sep)
                if last_sep != -1:
                    end = search_start + last_sep + len(sep)
                    break

        chunk = text[start:end].strip()
        if len(chunk) > 50:
            chunks.append(chunk)

        # Move start with overlap
        start = end - overlap if end < len(text) else len(text)

    return chunks


def generate_embeddings(texts: list[str], client: OpenAI) -> list[list[float]]:
    """Generate embeddings for texts."""
    if not texts:
        return []

    # Process in batches
    batch_size = 100
    all_embeddings = []

    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        response = client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=batch,
        )
        batch_embeddings = [item.embedding for item in response.data]
        all_embeddings.extend(batch_embeddings)

    return all_embeddings


def main():
    # Source directory
    source_dir = Path(
        "/Users/rbhatti/Documents/Medical_Device_Regulatory_Hub/Perplexity approach/All previous documents via perplexity  stage 0/Regulatory KB"
    )

    # Also include root level files
    root_dir = Path(
        "/Users/rbhatti/Documents/Medical_Device_Regulatory_Hub/Perplexity approach/All previous documents via perplexity  stage 0"
    )

    # Vector store directory
    vectorstore_dir = Path("/Users/rbhatti/health-canada-meddev-agent/data/vectorstore")
    vectorstore_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("Health Canada Regulatory Document Ingestion")
    print("=" * 60)

    # Initialize OpenAI client
    print("\n[1/5] Initializing OpenAI client...")
    client = OpenAI()

    # Initialize ChromaDB
    print("[2/5] Initializing ChromaDB...")
    chroma_client = chromadb.PersistentClient(
        path=str(vectorstore_dir),
        settings=ChromaSettings(anonymized_telemetry=False),
    )

    # Delete existing collection if exists, then create new
    try:
        chroma_client.delete_collection(COLLECTION_NAME)
    except Exception:
        pass

    collection = chroma_client.create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )

    # Find all documents
    print("[3/5] Finding documents...")
    documents = []

    # Search in Regulatory KB
    for ext in ["*.pdf", "*.docx", "*.md"]:
        documents.extend(source_dir.rglob(ext))

    # Also get root level files
    for ext in ["*.pdf", "*.docx", "*.md"]:
        documents.extend(root_dir.glob(ext))

    # Remove duplicates
    documents = list(set(documents))
    print(f"   Found {len(documents)} documents")

    # Process documents
    print("[4/5] Processing documents...")
    all_chunks = []
    all_metadatas = []
    all_ids = []

    stats = {"processed": 0, "failed": 0, "chunks": 0}

    for doc_path in documents:
        try:
            # Load document based on extension
            ext = doc_path.suffix.lower()
            if ext == ".pdf":
                doc_data = load_pdf(doc_path)
            elif ext in [".docx", ".doc"]:
                doc_data = load_docx(doc_path)
            elif ext in [".md", ".markdown"]:
                doc_data = load_markdown(doc_path)
            else:
                continue

            content = doc_data["content"]
            if not content or len(content) < 50:
                print(f"  Skipping (too short): {doc_path.name}")
                continue

            # Detect category
            category = detect_category(doc_path, content)

            # Chunk the content
            chunks = chunk_text(content)

            # Create entries for each chunk
            for i, chunk in enumerate(chunks):
                chunk_id = (
                    f"{doc_path.stem}_{hashlib.md5(chunk[:100].encode()).hexdigest()[:8]}_{i}"
                )

                metadata = {
                    "source": str(doc_path),
                    "file_name": doc_path.name,
                    "file_type": ext[1:],
                    "category": category,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                }

                all_chunks.append(chunk)
                all_metadatas.append(metadata)
                all_ids.append(chunk_id)

            stats["processed"] += 1
            stats["chunks"] += len(chunks)
            print(f"  Processed: {doc_path.name} ({len(chunks)} chunks)")

        except Exception as e:
            print(f"  ERROR processing {doc_path.name}: {e}")
            stats["failed"] += 1

    if not all_chunks:
        print("\nNo chunks to process!")
        return

    # Generate embeddings
    print(f"\n[5/5] Generating embeddings for {len(all_chunks)} chunks...")
    print("   This may take a few minutes...")

    embeddings = generate_embeddings(all_chunks, client)

    # Add to ChromaDB in batches
    print("   Storing in vector database...")
    batch_size = 100
    for i in range(0, len(all_chunks), batch_size):
        end = min(i + batch_size, len(all_chunks))
        collection.add(
            ids=all_ids[i:end],
            documents=all_chunks[i:end],
            embeddings=embeddings[i:end],
            metadatas=all_metadatas[i:end],
        )
        print(f"   Stored {end}/{len(all_chunks)} chunks")

    # Print summary
    print("\n" + "=" * 60)
    print("INGESTION COMPLETE")
    print("=" * 60)
    print(f"Documents processed: {stats['processed']}")
    print(f"Documents failed: {stats['failed']}")
    print(f"Total chunks: {stats['chunks']}")
    print(f"Vector store: {vectorstore_dir}")
    print(f"Collection: {COLLECTION_NAME}")

    # Show category breakdown
    categories = {}
    for meta in all_metadatas:
        cat = meta.get("category", "other")
        categories[cat] = categories.get(cat, 0) + 1

    print("\nChunks by category:")
    for cat, count in sorted(categories.items()):
        print(f"  {cat}: {count}")


if __name__ == "__main__":
    main()
